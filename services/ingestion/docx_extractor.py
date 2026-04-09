"""DOCX ingestion service for CloudInsight."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from services.ingestion.common import (
    IngestionValidationError,
    build_error_result,
    build_success_result,
    clean_text,
    count_words,
    join_text_fragments,
    validate_file,
)


logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract paragraphs and tables from DOCX files."""

    def __init__(self, preview_rows: int = 25) -> None:
        self.preview_rows = preview_rows

    def extract(self, file_path: str | Path) -> dict[str, Any]:
        """Extract text and table structures from a DOCX file."""
        try:
            path = validate_file(file_path, {".docx"})
        except IngestionValidationError as exc:
            return build_error_result(file_path, "docx", errors=[str(exc)])

        try:
            from docx import Document
        except ImportError:
            return build_error_result(
                path,
                "docx",
                errors=["Missing dependency 'python-docx'. Install it before ingesting DOCX files."],
            )

        try:
            document = Document(str(path))
            paragraph_sections: list[dict[str, Any]] = []
            table_sections: list[dict[str, Any]] = []
            tables: list[dict[str, Any]] = []

            for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
                paragraph_text = clean_text(paragraph.text)
                if not paragraph_text:
                    continue

                paragraph_sections.append(
                    {
                        "id": f"paragraph_{paragraph_number}",
                        "type": "paragraph",
                        "paragraph_number": paragraph_number,
                        "text": paragraph_text,
                        "char_count": len(paragraph_text),
                        "word_count": count_words(paragraph_text),
                    }
                )

            for table_number, table in enumerate(document.tables, start=1):
                preview_rows: list[dict[str, Any]] = []
                row_text_fragments: list[str] = []
                column_count = 0

                for row_number, row in enumerate(table.rows, start=1):
                    values = [clean_text(cell.text) for cell in row.cells]
                    column_count = max(column_count, len(values))
                    row_payload = {f"column_{index}": value for index, value in enumerate(values, start=1)}
                    preview_rows.append(row_payload)

                    flattened = " | ".join(value for value in values if value)
                    if flattened:
                        row_text_fragments.append(flattened)

                table_text = join_text_fragments(row_text_fragments)
                tables.append(
                    {
                        "id": f"table_{table_number}",
                        "name": f"Table {table_number}",
                        "row_count": len(preview_rows),
                        "column_count": column_count,
                        "preview_rows": preview_rows[: self.preview_rows],
                        "rows": preview_rows[: self.preview_rows],
                    }
                )

                if table_text:
                    table_sections.append(
                        {
                            "id": f"table_text_{table_number}",
                            "type": "table",
                            "table_number": table_number,
                            "text": table_text,
                            "char_count": len(table_text),
                            "word_count": count_words(table_text),
                        }
                    )

            sections = paragraph_sections + table_sections
            full_text = join_text_fragments(section["text"] for section in sections)
            warnings = []
            if not full_text:
                warnings.append("The DOCX file was parsed, but no extractable text content was found.")

            metadata = {
                "paragraph_count": len(paragraph_sections),
                "table_count": len(tables),
                "raw_paragraph_count": len(document.paragraphs),
            }

            logger.info("Extracted DOCX content from %s", path)
            return build_success_result(
                path,
                "docx",
                metadata=metadata,
                text=full_text,
                sections=sections,
                tables=tables,
                warnings=warnings,
            )
        except Exception as exc:
            logger.exception("DOCX extraction failed for %s", path)
            return build_error_result(path, "docx", errors=[f"DOCX extraction failed: {exc}"])


def extract_docx(file_path: str | Path, **kwargs: Any) -> dict[str, Any]:
    """Convenience wrapper for DOCX extraction."""
    return DOCXExtractor(**kwargs).extract(file_path)
